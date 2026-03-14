"""
Instagram運用ドキュメント2本をWord形式で生成する
- B: 現場運用マニュアル
- A: 職員会議向け提案資料
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "../docs/reports")

# ===== スタイルヘルパー =====

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_paragraph(doc, text="", bold=False, size=None, color=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダー行
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1E3A5F")
        tcPr.append(shd)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # データ行
    for r_i, row in enumerate(rows):
        row_cells = table.rows[r_i + 1].cells
        for c_i, cell_text in enumerate(row):
            row_cells[c_i].text = str(cell_text)
            row_cells[c_i].paragraphs[0].runs[0].font.size = Pt(9.5)
            if r_i % 2 == 1:
                tc = row_cells[c_i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EEF3F8")
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

def add_callout(doc, text, bg="F0F4F8"):
    """注意書きボックス（グレー背景段落で代替）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.italic = True
    # 背景色をシェーディングで
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0BEC5")
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "游ゴシック"
    style.font.size = Pt(10.5)
    # 見出しフォント設定
    for i in range(1, 5):
        try:
            h_style = doc.styles[f"Heading {i}"]
            h_style.font.name = "游ゴシック"
        except Exception:
            pass


# ===== B: 現場運用マニュアル =====

def build_field_manual():
    doc = Document()
    set_doc_defaults(doc)

    # マージン設定
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 表紙相当
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("鳥羽高校 Instagram 現場運用マニュアル")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("担当教員・顧問向け  |  企画推進部作成  |  2026年3月")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    add_callout(doc,
        "【対象読者】投稿担当の顧問・教諭、企画推進部担当\n"
        "【アカウント】@toba_hs_official\n"
        "【問い合わせ】企画推進部担当（内線___）",
        bg="E8F0FB"
    )

    doc.add_paragraph()

    # ===== 第0章 =====
    set_heading(doc, "0. このマニュアルの使い方", 1, (0x1E, 0x3A, 0x5F))
    add_bullet(doc, "初めて投稿する方 → 第1章〜第3章を通読してから実施")
    add_bullet(doc, "投稿を作成したい方 → 第2章「投稿の作り方」を確認")
    add_bullet(doc, "肖像権を確認したい方 → 第4章「顔出しOK台帳の使い方」を確認")
    add_bullet(doc, "緊急時（炎上・誤投稿） → 第6章「トラブル対応」を確認")

    add_separator(doc)

    # ===== 第1章 =====
    set_heading(doc, "第1章　全体の流れ（必ず読んでください）", 1, (0x1E, 0x3A, 0x5F))
    set_heading(doc, "1-1　投稿フロー（全担当者共通）", 2)

    flow_steps = [
        "① 撮影・コンテンツ準備",
        "↓",
        "② Meta Business Suite で「下書き」作成",
        "↓",
        "③ 企画推進部担当にグループチャットまたは口頭で連絡",
        "↓",
        "④ 管理職が内容確認・承認（修正指示がある場合あり）",
        "↓",
        "⑤ 管理職が公開（または日時指定）",
    ]
    for step in flow_steps:
        p = doc.add_paragraph(style="List Number" if step.startswith("①") or step.startswith("②") or step.startswith("③") or step.startswith("④") or step.startswith("⑤") else "Normal")
        p.clear()
        run = p.add_run(step)
        run.font.size = Pt(10.5)
        if step == "↓":
            run.font.color.rgb = RGBColor(150, 150, 150)
            p.paragraph_format.left_indent = Cm(1.5)

    add_callout(doc,
        "⚠ 重要: ステップ②で「今すぐ投稿」ボタンは絶対に押さないでください。\n"
        "担当教員の権限では「下書き保存まで」がルールです。",
        bg="FFF3CD"
    )

    set_heading(doc, "1-2　権限の種類と担当者", 2)
    add_table(doc,
        ["権限", "対象", "できること"],
        [
            ["管理者", "広報責任者・教頭", "全操作（公開・設定変更・アカウント管理）"],
            ["編集者", "各部活顧問・企画推進部担当", "下書き作成・編集・メッセージ返信"],
            ["モデレーター", "必要に応じ追加", "コメント管理・インサイト閲覧"],
        ],
        col_widths=[3, 5, 7]
    )

    add_separator(doc)

    # ===== 第2章 =====
    set_heading(doc, "第2章　投稿の作り方（Meta Business Suite 操作手順）", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "2-1　ログイン", 2)
    add_bullet(doc, "ブラウザで business.facebook.com にアクセス")
    add_bullet(doc, "学校共有アカウントの Facebook メールアドレス・パスワードでログイン")
    add_bullet(doc, "ログイン情報の管理: 企画推進部担当が保管。個人端末への保存は禁止。", level=1)
    add_bullet(doc, "左側メニューから「鳥羽高等学校」のビジネスアカウントを選択")
    add_callout(doc,
        "モバイルの場合: App Store / Google Play で「Meta Business Suite」をインストール。"
        "写真撮影後すぐに下書きを作りたい場合に便利。ただし最終確認はPCブラウザ版を推奨（プレビューが正確）。"
    )

    set_heading(doc, "2-2　下書きの作成手順（PC版）", 2)
    steps = [
        "左メニュー「投稿とストーリー」→「投稿を作成」をクリック",
        "Instagram を選択（Facebook との同時投稿は管理者に確認）",
        "写真・動画の選択\n　推奨サイズ: 正方形 1080×1080px、横長 1080×566px\n　動画（リール）: 縦型 1080×1920px、最大60秒（ストーリー）〜90秒（リール）",
        "キャプションの入力\n　学校名: 「京都府立鳥羽高等学校」を1行目に入れる（統一感が出る）\n　ハッシュタグ: 第3章を参照\n　個人名: 原則として姓のみ、またはイニシャル",
        "右上の「▼」から「下書きとして保存」を選択（「今すぐ投稿」は押さない）",
        "グループチャット（LINE / Teams 等）に「下書きを保存しました」と送信し、企画推進部担当に伝える",
    ]
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i+1}. {step}")
        run.font.size = Pt(10.5)

    set_heading(doc, "2-3　承認後の公開手順（管理者のみ）", 2)
    mgr_steps = [
        "「投稿とストーリー」→「コンテンツ」→「下書き」タブを開く",
        "該当投稿を選択し、内容を確認",
        "問題なければ「今すぐ投稿」または「日時指定」で公開",
        "修正が必要な場合は担当者にチャットで修正内容を伝え、再度下書きを確認する",
    ]
    for i, step in enumerate(mgr_steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{i+1}. {step}")
        run.font.size = Pt(10.5)

    add_separator(doc)

    # ===== 第3章 =====
    set_heading(doc, "第3章　コンテンツ基準とハッシュタグ", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "3-1　投稿してよいもの・ダメなもの", 2)
    add_paragraph(doc, "✅ 投稿OK", bold=True)
    add_table(doc,
        ["カテゴリ", "例"],
        [
            ["部活・大会", "大会出場報告、試合結果速報、練習風景（引き）"],
            ["行事・イベント", "文化祭、体育祭、校外学習、生徒会活動"],
            ["学習・探究", "プレゼンの様子（後ろ姿OK）、作品紹介、留学報告"],
            ["学校の日常", "食堂メニュー、図書館の本、校庭の季節の風景"],
            ["実績・受賞", "資格取得、コンクール入賞、進路実績（個人名なし）"],
        ],
        col_widths=[4, 11]
    )
    add_paragraph(doc, "❌ 投稿NG", bold=True)
    add_table(doc,
        ["NG事項", "理由"],
        [
            ["未同意の生徒の顔が写った写真", "肖像権・プライバシー侵害"],
            ["合否情報・成績・成績順位", "個人情報保護"],
            ["特定の生徒を過度に称賛・批判する内容", "人間関係トラブルの原因"],
            ["政治的・宗教的内容", "中立性の維持"],
            ["他校の批判・比較", "対外的トラブルの原因"],
        ],
        col_widths=[7, 8]
    )

    set_heading(doc, "3-2　ハッシュタグルール", 2)
    add_paragraph(doc, "基本セット（毎回付ける）", bold=True)
    add_callout(doc, "#鳥羽高校　#京都府立鳥羽高等学校　#京都府立高校　#鳥羽　#高校生", bg="E8F0FB")
    add_paragraph(doc, "カテゴリ別（投稿内容に合わせて追加）", bold=True)
    add_table(doc,
        ["カテゴリ", "ハッシュタグ例"],
        [
            ["部活（野球）", "#高校野球 #京都高校野球 #野球部"],
            ["部活（バスケ）", "#高校バスケ #バスケットボール部"],
            ["部活（ウエイトリフティング）", "#ウエイトリフティング #重量挙げ"],
            ["文化祭", "#文化祭 #高校文化祭"],
            ["留学・国際", "#海外研修 #グローバル教育"],
        ],
        col_widths=[5, 10]
    )
    add_callout(doc, "上限の目安: 合計10〜15個。多すぎるとスパム判定されることがある。")

    set_heading(doc, "3-3　キャプションテンプレート", 2)
    add_callout(doc,
        "【部活大会結果】\n"
        "○○部が○○大会に出場しました！\n"
        "\n"
        "[結果・コメント等を2〜3文で]\n"
        "\n"
        "引き続き応援よろしくお願いします！\n"
        "#鳥羽高校 #京都府立鳥羽高等学校 #京都府立高校 #高校生 #[部活名]",
        bg="F0F4F8"
    )

    add_separator(doc)

    # ===== 第4章 =====
    set_heading(doc, "第4章　肖像権確認手順（必読）", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "4-1　基本原則", 2)
    add_paragraph(doc,
        "入学時に「SNS・WEB肖像利用同意書」を全生徒・保護者から回収しています（在校生を含む全学年対象）。"
    )
    add_table(doc,
        ["同意状況", "顔の扱い"],
        [
            ["同意あり（全媒体またはSNSを含む）", "自然な表情で掲載可能"],
            ["同意あり（特定媒体のみ・SNS除外）", "顔を隠す加工が必要"],
            ["同意なし・不明", "スタンプ・ぼかし・後ろ姿等で顔を隠す"],
        ],
        col_widths=[7, 8]
    )
    add_callout(doc, "原則として「同意不明 = 同意なし」として扱ってください。", bg="FFF3CD")

    set_heading(doc, "4-2　肖像OK台帳の確認手順", 2)
    for i, step in enumerate([
        "企画推進部が管理する「肖像利用同意台帳」（Excel または印刷物）を参照",
        "写真に写っている生徒の名前・クラスを確認し、同意欄を照合",
        "1名でも「同意なし・不明」の生徒が写っている場合は、全員分の確認が済むまで投稿不可",
        "顔加工が必要な場合は、スマートフォンの標準写真編集機能またはLINEのスタンプ機能で対応",
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i+1}. {step}").font.size = Pt(10.5)

    set_heading(doc, "4-3　同意書が見つからない・不明な生徒への対応", 2)
    add_paragraph(doc,
        "万が一、台帳で同意状況が確認できない生徒が写真に含まれる場合は、"
        "「不明 = 同意なし」として扱い、顔を隠す加工を施してください。"
        "台帳の不備については企画推進部担当に報告してください。"
    )

    add_separator(doc)

    # ===== 第5章 =====
    set_heading(doc, "第5章　コメント・DM対応", 1, (0x1E, 0x3A, 0x5F))
    add_callout(doc, "基本方針: 原則として、コメントへの返信は行わない。", bg="FFF3CD")
    add_paragraph(doc,
        "公式アカウントへのコメントに逐一返信することで生まれる誤解・炎上リスクを避けるため、"
        "コメント返信は原則として行いません。"
    )
    set_heading(doc, "問い合わせへの返信テンプレート", 2)
    add_callout(doc,
        "ご質問ありがとうございます。\n"
        "詳しいお問い合わせは、公式ホームページのお問い合わせフォームをご利用ください。\n"
        "https://www.kyoto-be.ne.jp/toba-hs/",
        bg="E8F0FB"
    )
    set_heading(doc, "不適切コメントへの対応", 2)
    add_table(doc,
        ["内容", "対応"],
        [
            ["誹謗中傷・悪口", "非表示（アーカイブ）し、管理者に報告"],
            ["スパム・広告", "削除"],
            ["進路相談・個人的な問い合わせ", "公式HPへの誘導テンプレートで返信"],
        ],
        col_widths=[6, 9]
    )

    add_separator(doc)

    # ===== 第6章 =====
    set_heading(doc, "第6章　トラブル対応（緊急時）", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "6-1　誤投稿（間違えて公開してしまった）", 2)
    for i, step in enumerate([
        "即座に非表示（アーカイブ）にする → 投稿右上「…」→「アーカイブ」を選択",
        "管理者（広報責任者）にすぐ電話またはLINEで報告",
        "管理者が事実確認を行い、削除・修正・謝罪の要否を判断",
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i+1}. {step}").font.size = Pt(10.5)
    add_callout(doc,
        "アーカイブは外部からは見えなくなりますが、削除ではなく復元できます。"
        "まず落ち着いてアーカイブを選んでください。"
    )

    set_heading(doc, "6-2　炎上・問題視された投稿", 2)
    for i, step in enumerate([
        "管理者（教頭・校長）への即時報告",
        "管理者の指示があるまで投稿をアーカイブ（非表示）にする",
        "コメント欄をコメントオフ（投稿設定から変更可能）にする",
        "自己判断での謝罪コメントや返信は絶対にしない",
        "管理職主導で事実確認・公式見解を決定する",
    ]):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f"{i+1}. {step}").font.size = Pt(10.5)

    add_separator(doc)

    # ===== 第7章 =====
    set_heading(doc, "第7章　投稿の目安・頻度", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "7-1　推奨投稿ペース（導入初期）", 2)
    add_table(doc,
        ["期間", "目標頻度"],
        [
            ["導入〜3ヶ月", "週1〜2回（まず続けることを優先）"],
            ["4ヶ月〜6ヶ月", "週2〜3回（季節行事に合わせて増やす）"],
            ["安定期以降", "週3回以上（複数担当者で分担）"],
        ],
        col_widths=[5, 10]
    )
    add_callout(doc, "無理に毎日投稿する必要はありません。継続 > 頻度です。")

    set_heading(doc, "7-2　年間コンテンツカレンダー", 2)
    add_table(doc,
        ["時期", "コンテンツネタ"],
        [
            ["4月", "入学式、新入部員歓迎、春の部活動開始"],
            ["5月", "中間考査前の様子、春季大会"],
            ["6月", "文化祭準備、体育祭、学校見学会告知"],
            ["7月", "夏季大会（野球・バスケ等）、オープンスクール"],
            ["8月", "夏合宿、インターハイ、大会速報"],
            ["9月", "秋の行事（遠足・修学旅行）、後期開始"],
            ["10月", "文化祭本番、学校見学会"],
            ["11月", "進路報告（指定校推薦など）、部活秋季大会"],
            ["12月", "冬季大会、年間振り返り投稿"],
            ["1月", "共通テスト激励・応援投稿（個人名なし）"],
            ["2月", "前期選抜告知、部活コンテンツ"],
            ["3月", "卒業式、合格発表（速報・アカウント紹介）"],
        ],
        col_widths=[3, 12]
    )

    add_separator(doc)

    # ===== 第8章 Q&A =====
    set_heading(doc, "第8章　よくある質問（Q&A）", 1, (0x1E, 0x3A, 0x5F))

    qas = [
        ("自分のスマートフォンで撮った写真を使っていいですか？",
         "使えます。ただし学校アカウントのコンテンツとして管理するため、写真はグループチャットや共有ドライブを経由して企画推進部に提出してください。個人端末から直接 Meta Business Suite に投稿しないよう注意してください。"),
        ("生徒が「自分の写真を載せてほしい」と言ってきました。",
         "生徒本人の希望があっても、保護者（未成年の場合）の同意確認が必要です。同意台帳で確認するか、保護者への確認を経てから対応してください。"),
        ("試合の速報をリアルタイムで投稿したいのですが、承認待ちが間に合いません。",
         "管理者（教頭等）にあらかじめ「○○大会に出場するので即日速報投稿を想定しています」と伝え、事前承認を取ってください。「試合結果と選手の後ろ姿写真のみ」であれば事前包括承認で顧問が直接投稿することも可能です。管理者と事前に合意してください。"),
        ("間違えて「今すぐ投稿」を押してしまいました。",
         "すぐにアーカイブにして管理者に報告してください（第6章参照）。焦って別の操作をするより、まずアーカイブが最善です。"),
    ]

    for q, a in qas:
        p_q = doc.add_paragraph()
        r = p_q.add_run(f"Q. {q}")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Cm(0.5)
        r_a = p_a.add_run(f"A. {a}")
        r_a.font.size = Pt(10.5)
        doc.add_paragraph()

    add_separator(doc)

    # ===== 付録C: チェックリスト =====
    set_heading(doc, "付録　投稿前チェックリスト", 1, (0x1E, 0x3A, 0x5F))
    checklist = [
        "写真に写っている生徒全員の肖像権を確認した（台帳照合またはぼかし処理済み）",
        "個人名（フルネーム）を使用していない",
        "学校名を含む基本ハッシュタグを入れた",
        "キャプションに誤字・脱字がない",
        "「今すぐ投稿」ではなく「下書き保存」を選んだ",
        "企画推進部担当に下書き完成を連絡した",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"□  {item}")
        run.font.size = Pt(10.5)

    doc.add_paragraph()
    add_callout(doc, "本マニュアルの内容について不明点・改善提案は企画推進部担当まで。定期的に（年1回以上）見直します。")

    out_path = os.path.join(OUTPUT_DIR, "instagram_field_manual.docx")
    doc.save(out_path)
    print(f"保存: {out_path}")
    return out_path


# ===== A: 職員会議向け提案資料 =====

def build_proposal():
    doc = Document()
    set_doc_defaults(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 表紙
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Instagram 公式アカウント開設・運用開始 提案書")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("職員会議提出用  |  企画推進部  |  2026年○月○日")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # 1. 提案概要
    set_heading(doc, "1. 提案の概要", 1, (0x1E, 0x3A, 0x5F))
    add_paragraph(doc,
        "本校の中学生向け広報強化を目的として、"
        "学校公式Instagramアカウントの開設と運用体制の整備を提案します。"
    )
    add_table(doc,
        ["項目", "内容"],
        [
            ["提案内容", "公式Instagramアカウントの開設・運用ルールの制定"],
            ["目的", "中学生・保護者への認知拡大、第一志望校としての選択率向上"],
            ["費用", "原則無料（Meta Business Suite の標準機能を使用）"],
            ["運用開始目標", "2026年○月"],
        ],
        col_widths=[4, 11]
    )

    add_separator(doc)

    # 2. なぜ今
    set_heading(doc, "2. なぜ今、Instagramなのか", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "2-1. 府立高校の導入状況（2026年3月調査）", 2)
    add_callout(doc,
        "京都府立高校46校のうち、35校（76.1%）がInstagramを公式運用中。\n"
        "本校（鳥羽高校）は現時点で未開設であり、府立高校の中では少数派に属します。",
        bg="FFF3CD"
    )
    add_paragraph(doc,
        "競合する周辺の府立高校（西京・堀川・紫野・桂など）もすでに運用を開始しています。"
        "Instagramは「見つけてもらうための入口」として、高校広報の標準インフラになっています。"
    )

    set_heading(doc, "2-2. 中学生の情報収集行動の変化", 2)
    add_bullet(doc, "中学生・保護者が高校を調べる際、最初に見るのはGoogle検索やSNSであり、学校HPは後半に確認される傾向があります")
    add_bullet(doc, "Instagramは「学校の雰囲気・部活の熱量・日常の様子」を一目で伝えられる唯一のメディアです")
    add_bullet(doc, "現在の本校HPは更新頻度が低く、リアルタイムの情報発信が困難な構造です")

    set_heading(doc, "2-3. 鳥羽高校が発信すべき強み", 2)
    add_table(doc,
        ["強み", "具体的な内容"],
        [
            ["野球部", "甲子園常連クラスの競争力"],
            ["バスケットボール部・ウエイトリフティング部", "全国レベルの実績"],
            ["グローバル科", "海外研修・留学プログラム"],
            ["立地の希少性", "京都駅から最も近い公立高校"],
            ["探究学習・進路実績", "多様な学習環境と進路指導"],
        ],
        col_widths=[6, 9]
    )

    add_separator(doc)

    # 3. 運用体制
    set_heading(doc, "3. 運用体制と管理職の関与", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "3-1. 承認フロー（「管理職承認なしには公開されない」仕組み）", 2)
    add_callout(doc,
        "全投稿は管理職の確認・承認を経てから公開されます。担当教員が勝手に公開できない設計です。\n\n"
        "①担当教員が「下書き」を作成\n"
        "②企画推進部にチャットで連絡\n"
        "③管理職（教頭等）が内容を確認・承認\n"
        "④管理職が「公開」ボタンを押す\n\n"
        "担当教員には「下書き作成」の権限のみ付与します。「公開」は管理者権限を持つ方のみです。\n"
        "この仕組みは Meta Business Suite（Facebookが提供する無料の公式管理ツール）で実現します。",
        bg="E8F0FB"
    )

    set_heading(doc, "3-2. 運用責任者と担当者", 2)
    add_table(doc,
        ["役割", "担当"],
        [
            ["運用責任者（管理者権限）", "企画推進部長、教頭"],
            ["投稿担当（下書き作成）", "各部活顧問、企画推進部担当教員"],
            ["ガイドライン管理", "企画推進部"],
        ],
        col_widths=[6, 9]
    )

    add_separator(doc)

    # 4. 個人情報
    set_heading(doc, "4. 個人情報・肖像権の取り扱い", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "4-1. 生徒の肖像利用ルール", 2)
    add_paragraph(doc,
        "入学時に保護者から「SNS・WEB肖像利用同意書」を回収しています（在校生を含む全学年対象）。"
    )
    add_table(doc,
        ["同意状況", "取り扱い"],
        [
            ["同意あり", "自然な表情で掲載可"],
            ["同意なし・不明", "顔を隠す加工（スタンプ・ぼかし等）を施す"],
        ],
        col_widths=[6, 9]
    )
    add_callout(doc, "「同意不明 = 同意なし」の原則を徹底します。", bg="FFF3CD")

    set_heading(doc, "4-2. 投稿禁止事項（要点）", 2)
    for item in [
        "生徒の合否・成績等に関わる情報",
        "個人を特定できる住所等",
        "政治的・宗教的内容",
        "他校への批判・比較",
    ]:
        add_bullet(doc, item)

    add_separator(doc)

    # 5. リスク管理
    set_heading(doc, "5. リスク管理", 1, (0x1E, 0x3A, 0x5F))

    set_heading(doc, "5-1. 炎上・誤投稿への対応", 2)
    add_table(doc,
        ["事態", "初動対応"],
        [
            ["誤投稿した場合", "即座に非表示（アーカイブ）→ 管理者に報告"],
            ["炎上・問題視された場合", "即座に非表示 → 管理職主導で公式見解を決定"],
            ["不正ログイン", "パスワード変更・二要素認証再設定"],
        ],
        col_widths=[6, 9]
    )
    add_callout(doc,
        "「アーカイブ」は公開を一時停止するだけで、削除ではありません。復元できます。\n"
        "誤投稿が起きた際も、落ち着いて手順通りに対応できる体制を整えます。"
    )

    set_heading(doc, "5-2. コメント・DM対応方針", 2)
    add_bullet(doc, "原則としてコメントへの返信は行いません（誤解・炎上リスクを回避）")
    add_bullet(doc, "問い合わせは「公式HPのお問い合わせフォームへ誘導」するテンプレートを使用します")

    add_separator(doc)

    # 6. 費用
    set_heading(doc, "6. 費用", 1, (0x1E, 0x3A, 0x5F))
    add_table(doc,
        ["項目", "費用"],
        [
            ["Instagramアカウント開設", "無料"],
            ["Meta Business Suite（承認ワークフロー）", "無料"],
            ["機器・設備", "既存のPC・スマートフォンを使用"],
            ["将来の有料ツール移行", "投稿頻度増加後に検討（月額数千円〜）"],
        ],
        col_widths=[8, 7]
    )
    add_callout(doc, "導入初期の追加費用はありません。", bg="E8F5E9")

    add_separator(doc)

    # 7. スケジュール
    set_heading(doc, "7. 導入スケジュール（案）", 1, (0x1E, 0x3A, 0x5F))
    add_table(doc,
        ["時期", "マイルストーン"],
        [
            ["承認後〜1ヶ月", "アカウント開設・権限設定・同意書様式確定"],
            ["〜2ヶ月", "担当教員へのマニュアル共有・操作研修（30分程度）"],
            ["〜3ヶ月", "試験運用開始（週1〜2投稿、管理者確認を徹底）"],
            ["6ヶ月後", "運用実績をもとに頻度・担当体制を見直し"],
        ],
        col_widths=[4, 11]
    )

    add_separator(doc)

    # 8. 承認事項
    set_heading(doc, "8. 承認をお願いしたいこと", 1, (0x1E, 0x3A, 0x5F))
    add_callout(doc,
        "以下2点について承認をお願いします。\n\n"
        "1. 公式Instagramアカウントの開設\n"
        "2. 「Instagram現場運用マニュアル」に基づく運用体制の整備・運用開始",
        bg="E8F0FB"
    )

    doc.add_paragraph()
    add_paragraph(doc, "参考資料", bold=True)
    for ref in [
        "B: Instagram 現場運用マニュアル（別紙）",
        "Meta Business Suite 技術調査報告書",
        "Webプレゼン資料（府内競合校の分析と鳥羽高校の広報戦略全体像）",
    ]:
        add_bullet(doc, ref)

    out_path = os.path.join(OUTPUT_DIR, "instagram_proposal_for_staff_meeting.docx")
    doc.save(out_path)
    print(f"保存: {out_path}")
    return out_path


# ===== A-Simple: 職員会議向け提案資料（A4一枚版） =====

PRESENTATION_URL = "https://kerikerijy.github.io/TobaBuzzPR/presentation/index.html"

def build_proposal_one_pager():
    doc = Document()
    set_doc_defaults(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Instagram 公式アカウント開設・運用開始のご提案")
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("企画推進部  ｜  2026年○月○日").font.size = Pt(9)
    p2.paragraph_format.space_after = Pt(6)

    add_separator(doc)

    # ■ 背景・目的
    set_heading(doc, "■ 背景と目的", 2)
    p_bg = doc.add_paragraph()
    p_bg.paragraph_format.space_after = Pt(2)
    p_bg.add_run(
        "京都府立高校46校のうち35校（76.1%）がInstagramを公式運用しており、"
        "本校は未開設の少数派です。中学生・保護者の情報収集はSNSが起点となっており、"
        "このまま未開設を続けることは広報機会の損失につながります。\n"
        "詳細な競合分析・戦略背景については下記プレゼン資料をご覧ください。"
    ).font.size = Pt(10)

    # プレゼン資料リンク
    p_link = doc.add_paragraph()
    p_link.paragraph_format.left_indent = Cm(0.5)
    p_link.paragraph_format.space_after = Pt(4)
    run_label = p_link.add_run("▶ 広報戦略プレゼン資料: ")
    run_label.font.size = Pt(9.5)
    run_label.bold = True
    run_url = p_link.add_run(PRESENTATION_URL)
    run_url.font.size = Pt(9.5)
    run_url.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    # ハイパーリンクのXML追加
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    r_id = doc.part.relate_to(PRESENTATION_URL, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    rPr = run_url._r.get_or_add_rPr()
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.insert(0, rStyle)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.append(run_url._r)
    p_link._p.append(hyperlink)

    # ■ 運用体制
    set_heading(doc, "■ 運用体制（管理職承認なしには公開されない仕組み）", 2)
    add_table(doc,
        ["ステップ", "担当", "内容"],
        [
            ["①作成", "顧問・担当教員", "Meta Business Suite で下書き保存"],
            ["②連絡", "顧問・担当教員", "グループチャットで企画推進部に通知"],
            ["③確認", "管理職（教頭等）", "内容確認・修正指示"],
            ["④公開", "管理職", "承認後に「公開」ボタンを押す"],
        ],
        col_widths=[2.5, 4, 8.5]
    )

    # ■ 肖像権・個人情報
    set_heading(doc, "■ 個人情報・肖像権の取り扱い", 2)
    p_pr = doc.add_paragraph()
    p_pr.paragraph_format.space_after = Pt(2)
    p_pr.add_run(
        "入学時に取得済みの「SNS・WEB肖像利用同意書」をもとに掲載可否を管理します。"
        "「同意不明＝同意なし」の原則を徹底し、未同意の生徒の顔は必ず加工します。"
        "合否・成績等の個人情報、政治的・宗教的内容の投稿は禁止します。"
    ).font.size = Pt(10)

    # ■ リスク管理
    set_heading(doc, "■ リスク管理", 2)
    add_table(doc,
        ["事態", "初動対応"],
        [
            ["誤投稿・炎上", "即アーカイブ（非公開）→ 管理職へ報告 → 管理職主導で対応"],
            ["コメント・DM", "原則返信なし。問い合わせは公式HPフォームへ誘導"],
        ],
        col_widths=[4, 11]
    )

    # ■ 費用・スケジュール
    set_heading(doc, "■ 費用・スケジュール", 2)
    add_table(doc,
        ["項目", "内容"],
        [
            ["費用", "初期費用ゼロ（Meta Business Suite は無料）"],
            ["承認後〜1ヶ月", "アカウント開設・権限設定"],
            ["〜2ヶ月", "担当教員マニュアル共有・操作研修（30分）"],
            ["〜3ヶ月", "試験運用開始（週1〜2投稿）"],
        ],
        col_widths=[4, 11]
    )

    # ■ 承認事項
    add_separator(doc)
    set_heading(doc, "■ 承認事項", 2)
    add_callout(doc,
        "1. 公式Instagramアカウント（@toba_hs_official）の開設\n"
        "2. 「Instagram現場運用マニュアル」に基づく運用体制の整備・運用開始",
        bg="E8F0FB"
    )

    out_path = os.path.join(OUTPUT_DIR, "instagram_proposal_one_pager.docx")
    doc.save(out_path)
    print(f"保存: {out_path}")
    return out_path


if __name__ == "__main__":
    build_field_manual()
    build_proposal()
    build_proposal_one_pager()
    print("完了")
